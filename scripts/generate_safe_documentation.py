from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "SAFE_Documentacao_Completa.docx"
LOGO = ROOT / "public" / "images" / "logo.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_GRAY)
        set_cell_text(table.rows[0].cells[idx], header, bold=True)

    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)

    doc.add_paragraph()


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.style = "Code Block"
    p.paragraph_format.keep_together = True
    for line in code.strip().splitlines():
        run = p.add_run(line.rstrip())
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(31, 31, 31)
        p.add_run("\n")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code_style = doc.styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.2)
    code_style.paragraph_format.right_indent = Inches(0.2)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.0

    footer = section.footer.paragraphs[0]
    footer.text = "SAFE - Sistema de Autorização e Fluxo Escolar"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(89, 89, 89)


def add_cover(doc: Document) -> None:
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.35))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("SAFE")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Sistema de Autorização e Fluxo Escolar")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.add_run(
        "Documentação completa para criação, configuração, desenvolvimento, testes e apresentação acadêmica da aplicação."
    )

    add_table(
        doc,
        ["Item", "Descrição"],
        [
            ["Projeto", "Painel Laravel/Filament para autorizações escolares digitais"],
            ["Perfis", "Admin, AQV, Professores e Portaria"],
            ["Banco", "MySQL em desenvolvimento e SQLite em testes"],
            ["Data", "Maio de 2026"],
        ],
        [1.6, 4.7],
    )
    doc.add_page_break()


def add_overview(doc: Document) -> None:
    add_heading(doc, "1. Visão geral da aplicação")
    add_para(
        doc,
        "O SAFE foi criado para substituir autorizações em papel por um fluxo digital rastreável. A aplicação centraliza alunos, turmas, cursos, funcionários, autorizações, mensagens e notificações internas no painel administrativo.",
    )
    add_bullets(
        doc,
        [
            "AQV registra e acompanha solicitações de entrada e saída.",
            "Professores recebem autorizações das turmas vinculadas e registram leitura, aprovação, recusa ou observações.",
            "Portaria acompanha liberações e confirma saídas quando necessário.",
            "Admin gerencia funcionários, permissões, cursos, turmas e dados operacionais.",
        ],
    )

    add_heading(doc, "1.1 Problema resolvido", 2)
    add_para(
        doc,
        "Em processos manuais, autorizações podem se perder, atrasar ou não deixar evidências suficientes sobre quem solicitou, validou e finalizou a movimentação do aluno. O SAFE organiza esse fluxo em etapas digitais, com auditoria e notificações persistidas.",
    )

    add_heading(doc, "1.2 Módulos principais", 2)
    add_table(
        doc,
        ["Módulo", "Responsabilidade"],
        [
            ["Funcionários", "CRUD administrativo de usuários dos setores AQV, Portaria, Professores e Admin."],
            ["Cursos e turmas", "Organização acadêmica e vínculo de professores a múltiplas turmas."],
            ["Alunos", "Cadastro dos estudantes e associação com turmas."],
            ["Autorizações", "Entrada tardia e saída antecipada com leitura, aprovação, recusa e encerramento."],
            ["Notificações", "Alertas internos persistidos com status de leitura e badges."],
            ["Mensagens", "Comunicação interna por usuário ou setor."],
            ["Históricos", "Auditoria de ações de autorizações."],
        ],
        [1.8, 4.5],
    )


def add_setup(doc: Document) -> None:
    add_heading(doc, "2. Setup inicial")
    add_heading(doc, "2.1 PHP", 2)
    add_para(doc, "Instale PHP 8.3 ou superior. Em ambientes Windows com Laragon, selecione a versão PHP 8.3 no menu do Laragon.")
    add_code(doc, "php -v")

    add_heading(doc, "2.2 Composer", 2)
    add_para(doc, "Composer é o gerenciador de dependências PHP usado para instalar Laravel, Filament e pacotes auxiliares.")
    add_code(doc, "composer -V")

    add_heading(doc, "2.3 Node.js e NPM", 2)
    add_para(doc, "Node e NPM são usados para compilar CSS/JavaScript via Vite.")
    add_code(doc, "node -v\nnpm -v")

    add_heading(doc, "2.4 Banco de dados", 2)
    add_para(doc, "O projeto foi preparado para MySQL em desenvolvimento local e SQLite em memória nos testes. Para MySQL, crie um banco chamado safe.")
    add_code(doc, "CREATE DATABASE safe CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;")

    add_heading(doc, "2.5 Variáveis de ambiente", 2)
    add_code(
        doc,
        """
cp .env.example .env
php artisan key:generate
""",
    )
    add_para(doc, "Exemplo de configuração local MySQL:")
    add_code(
        doc,
        """
DB_CONNECTION=mysql
DB_HOST=127.0.0.1
DB_PORT=3308
DB_DATABASE=safe
DB_USERNAME=root
DB_PASSWORD=
""",
    )


def add_project_creation(doc: Document) -> None:
    add_heading(doc, "3. Criação do projeto do zero")
    add_heading(doc, "3.1 Criar Laravel", 2)
    add_code(doc, "composer create-project laravel/laravel SAFE\ncd SAFE")

    add_heading(doc, "3.2 Instalar autenticação e painel", 2)
    add_code(
        doc,
        """
composer require laravel/breeze --dev
php artisan breeze:install blade
composer require filament/filament spatie/laravel-permission
php artisan filament:install --panels
""",
    )

    add_heading(doc, "3.3 Dependências front-end", 2)
    add_code(doc, "npm install\nnpm run build")

    add_heading(doc, "3.4 Configuração de idioma", 2)
    add_para(
        doc,
        "Defina APP_LOCALE, APP_FALLBACK_LOCALE e APP_FAKER_LOCALE como pt_BR no arquivo .env. Também é recomendável manter APP_TIMEZONE como America/Sao_Paulo.",
    )

    add_heading(doc, "3.5 Configuração do painel Filament", 2)
    add_para(
        doc,
        "O painel principal fica em /admin. A classe AdminPanelProvider define nome SAFE, logo, tema, grupos de navegação, login customizado, recursos descobertos automaticamente e notificações de banco com polling.",
    )


def add_database(doc: Document) -> None:
    add_heading(doc, "4. Banco de dados, migrations e relacionamentos")
    add_table(
        doc,
        ["Tabela", "Campos principais", "Relacionamentos"],
        [
            ["users", "name, email, password, sector, is_active", "roles Spatie; muitas turmas como professor"],
            ["courses", "name, slug, description", "tem muitas turmas"],
            ["classrooms", "name, course_id, course, teacher_id", "pertence a curso; tem alunos; muitos professores"],
            ["classroom_teacher", "classroom_id, user_id", "pivot entre turmas e professores"],
            ["students", "name, registration, classroom_id", "pertence a turma; tem autorizações"],
            ["authorization_types", "name", "Entrada ou Saída"],
            ["authorizations", "student_id, type, status, reason, event_at, responsável, observações", "pertence a aluno, tipo, solicitante e processador"],
            ["authorization_audits", "authorization_id, user_id, action, note", "histórico de ações"],
            ["notifications", "type, notifiable, data, read_at", "notificações persistidas Laravel"],
            ["internal_messages", "sender, recipient, recipient_role, subject, body, read_at", "mensagens internas"],
        ],
        [1.65, 2.4, 2.25],
    )

    add_heading(doc, "4.1 Relacionamento professor-turma", 2)
    add_para(
        doc,
        "O SAFE usa uma relação muitos-para-muitos entre professores e turmas. Essa decisão permite que Eduardo acesse DS e Eletroeletrônica ao mesmo tempo em que Samuel fica limitado à Eletroeletrônica e Bruno ao Desenvolvimento de Sistemas.",
    )

    add_heading(doc, "4.2 Comandos de banco", 2)
    add_code(doc, "php artisan migrate\nphp artisan db:seed\n# Para recriar ambiente de demonstração\nphp artisan migrate:fresh --seed")


def add_models_and_seeders(doc: Document) -> None:
    add_heading(doc, "5. Models, seeders e dados de demonstração")
    add_heading(doc, "5.1 Models", 2)
    add_bullets(
        doc,
        [
            "User implementa FilamentUser para controlar acesso ao painel com canAccessPanel.",
            "User possui teachingClassrooms para o vínculo muitos-para-muitos com turmas.",
            "Classroom pertence a Course, possui Students e Teachers.",
            "Authorization concentra os métodos markAsRead, approve, deny e finish.",
            "InternalMessage notifica destinatários ao ser criada.",
        ],
    )

    add_heading(doc, "5.2 Seeders obrigatórios", 2)
    add_para(doc, "Os seeders criam automaticamente cursos, turmas, alunos, tipos de autorização, perfis e usuários de teste.")
    add_table(
        doc,
        ["Professor", "E-mail", "Escopo de acesso"],
        [
            ["Eduardo", "eduardo@safe.com", "Desenvolvimento de Sistemas e Eletroeletrônica"],
            ["Samuel", "samuel@safe.com", "Somente Eletroeletrônica"],
            ["Bruno", "bruno@safe.com", "Somente Desenvolvimento de Sistemas"],
        ],
        [1.4, 2.0, 2.9],
    )
    add_para(doc, "Cada um dos dois cursos recebe três turmas, e cada turma recebe cinco alunos com nomes e matrículas realistas.")


def add_permissions(doc: Document) -> None:
    add_heading(doc, "6. Autenticação, permissões e policies")
    add_para(doc, "A autorização combina Spatie Laravel Permission, policies do Laravel e filtros de consulta nos Resources do Filament.")
    add_table(
        doc,
        ["Perfil", "Acesso principal"],
        [
            ["Admin", "Acesso completo, incluindo CRUD de funcionários e exclusões administrativas."],
            ["AQV", "Gerencia autorizações, cursos, turmas e alunos."],
            ["Professor", "Visualiza alunos, turmas e autorizações apenas das turmas vinculadas."],
            ["Portaria", "Acompanha autorizações, confirma saídas e usa mensagens/notificações."],
        ],
        [1.5, 4.8],
    )
    add_heading(doc, "6.1 Policies implementadas", 2)
    add_bullets(
        doc,
        [
            "UserPolicy limita gestão de funcionários ao Admin.",
            "StudentPolicy e ClassroomPolicy aplicam escopo por turma.",
            "AuthorizationPolicy permite visualização por setor e restringe edição estrutural a Admin/AQV.",
            "InternalMessagePolicy limita visualização a remetente, destinatário, setor ou Admin.",
            "SystemNotificationPolicy garante que cada usuário veja apenas suas notificações.",
        ],
    )


def add_filament(doc: Document) -> None:
    add_heading(doc, "7. Filament e front-end administrativo")
    add_para(
        doc,
        "O front-end principal é o painel Filament. A navegação é organizada em Operação escolar, Comunicação, Cadastros acadêmicos e Administração.",
    )
    add_table(
        doc,
        ["Resource", "Tela"],
        [
            ["Employees", "Funcionários, setor, status ativo e turmas vinculadas ao professor."],
            ["Courses", "Cursos técnicos."],
            ["Classrooms", "Turmas, curso, professores vinculados e quantidade de alunos."],
            ["Students", "Alunos, matrícula, turma e curso."],
            ["Authorizations", "Autorizações digitais e ações do fluxo."],
            ["Notifications", "Notificações persistidas e status de leitura."],
            ["InternalMessages", "Mensagens por destinatário ou setor."],
            ["AuthorizationAudits", "Históricos somente leitura."],
        ],
        [2.2, 4.1],
    )
    add_heading(doc, "7.1 UX/UI", 2)
    add_bullets(
        doc,
        [
            "Login customizado com identidade SAFE e atalhos de usuários de teste.",
            "Badges de notificações e mensagens não lidas.",
            "Filtros por status, tipo de autorização e turma.",
            "Ações de fluxo diretamente na tabela para reduzir cliques.",
            "Campos de observação separados para AQV, professor e portaria.",
        ],
    )


def add_business_flows(doc: Document) -> None:
    add_heading(doc, "8. Regras de negócio e fluxos digitais")
    add_heading(doc, "8.1 Entrada de aluno atrasado", 2)
    add_numbered(
        doc,
        [
            "Portaria identifica ou encaminha o aluno atrasado.",
            "AQV cria autorização de Entrada com aluno, horário, motivo, responsável e observações.",
            "Professores vinculados à turma recebem notificação.",
            "Professor confirma leitura e registra observação se necessário.",
            "Professor aprova ou recusa a autorização.",
            "Sistema grava status, usuário responsável e auditoria.",
        ],
    )

    add_heading(doc, "8.2 Saída antecipada", 2)
    add_numbered(
        doc,
        [
            "AQV cria autorização do tipo Saída.",
            "Professor recebe notificação por pertencer à turma do aluno.",
            "Professor confirma leitura e aprova/libera ou recusa.",
            "Quando aprovado, a portaria recebe notificação de liberação.",
            "Portaria confirma a saída e finaliza o fluxo.",
            "Sistema registra horários, responsáveis, status e histórico completo.",
        ],
    )

    add_heading(doc, "8.3 Status de autorização", 2)
    add_table(
        doc,
        ["Status", "Significado"],
        [
            ["pending", "Aguardando validação."],
            ["approved", "Aprovada ou liberada pelo professor/AQV."],
            ["denied", "Recusada com justificativa."],
            ["finished", "Fluxo encerrado pela AQV ou portaria."],
        ],
        [1.4, 4.9],
    )


def add_notifications_and_messages(doc: Document) -> None:
    add_heading(doc, "9. Notificações, mensagens e eventos internos")
    add_para(
        doc,
        "As notificações usam a tabela notifications do Laravel. Cada notificação possui payload JSON com título, mensagem, URL, categoria e cor. O painel Filament exibe notificações de banco com polling configurado.",
    )
    add_bullets(
        doc,
        [
            "Criação de autorização notifica professores da turma.",
            "Leitura notifica o solicitante.",
            "Aprovação de saída notifica a portaria.",
            "Recusa notifica solicitante e professores envolvidos.",
            "Finalização notifica solicitante e professores.",
            "Mensagem interna notifica destinatário direto ou setor.",
        ],
    )

    add_heading(doc, "9.1 Serviço SafeNotifier", 2)
    add_para(
        doc,
        "SafeNotifier centraliza as regras de destinatários. Essa camada evita duplicação de lógica nos models e nas telas, e facilita a evolução futura para broadcast em tempo real.",
    )


def add_testing_and_deploy(doc: Document) -> None:
    add_heading(doc, "10. Testes, validação e deploy")
    add_heading(doc, "10.1 Testes automatizados", 2)
    add_para(doc, "A suíte cobre autenticação, seeders, escopo de professores, fluxos de autorização, mensagens, notificações e renderização das principais telas do painel.")
    add_code(doc, "php artisan test")

    add_heading(doc, "10.2 Qualidade de código", 2)
    add_code(doc, "vendor/bin/pint --test\nnpm run build")

    add_heading(doc, "10.3 Checklist de validação", 2)
    add_bullets(
        doc,
        [
            "Executar migrate:fresh --seed sem erros.",
            "Confirmar login com Admin, AQV, Portaria e professores.",
            "Validar que Eduardo vê DS e Eletroeletrônica.",
            "Validar que Samuel vê apenas Eletroeletrônica.",
            "Validar que Bruno vê apenas Desenvolvimento de Sistemas.",
            "Criar autorização de Entrada e confirmar leitura/aprovação.",
            "Criar autorização de Saída e finalizar pela portaria.",
            "Enviar mensagem interna e marcar como lida.",
            "Verificar badges e listagens de notificações.",
        ],
    )

    add_heading(doc, "10.4 Deploy", 2)
    add_numbered(
        doc,
        [
            "Configurar servidor com PHP 8.3, extensões PHP, Composer, Node e banco MySQL.",
            "Clonar o repositório e configurar .env de produção.",
            "Executar composer install --no-dev --optimize-autoloader.",
            "Executar npm ci e npm run build.",
            "Executar php artisan migrate --force.",
            "Configurar storage:link, permissões de storage e cache.",
            "Executar php artisan config:cache, route:cache e view:cache.",
            "Apontar o servidor web para public/.",
        ],
    )


def add_appendices(doc: Document) -> None:
    add_heading(doc, "11. Apêndice: comandos de referência")
    add_code(
        doc,
        """
composer install
npm install
cp .env.example .env
php artisan key:generate
php artisan migrate:fresh --seed
npm run build
php artisan serve
php artisan test
vendor/bin/pint --test
""",
    )

    add_heading(doc, "12. Melhorias futuras")
    add_bullets(
        doc,
        [
            "Substituir polling por WebSockets com Laravel Reverb ou serviço compatível.",
            "Criar relatórios por curso, turma, período, motivo e setor.",
            "Adicionar QR Code para conferência de autorização na portaria.",
            "Adicionar cadastro de responsáveis legais e contato de emergência.",
            "Exportar autorizações em PDF.",
            "Criar dashboard individual para AQV, professores e portaria.",
            "Integrar com sistemas acadêmicos externos.",
        ],
    )


def build() -> None:
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_overview(doc)
    add_setup(doc)
    add_project_creation(doc)
    add_database(doc)
    add_models_and_seeders(doc)
    add_permissions(doc)
    add_filament(doc)
    add_business_flows(doc)
    add_notifications_and_messages(doc)
    add_testing_and_deploy(doc)
    add_appendices(doc)

    for section in doc.sections:
        section.start_type = WD_SECTION.CONTINUOUS

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
